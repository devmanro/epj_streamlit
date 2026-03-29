import streamlit as st
import re
import os
import pandas as pd
from pyarrow import null
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from assets.constants.constants import (
    COL_CLIENT,
    COL_QUANTITE,
    COL_TONAGE,
    COL_BL,
    DB_PATH, COLUMNS,
    # add COL_VALUES here if you have such a column name
    COL_TYPE,
    COMMODITY_TYPES,
    UNITS_TYPES,
    PACKAGES_TYPES,
    GOODS__TYPES, COL_DESIGNATION,
    KEYWORD_RULES,
    numeric_cols,
    date_cols,
    category_cols,
    text_cols,
)




def clean_dataframe_types(datasource , only_cols=None):
    """
    Loops through columns and applies appropriate types and handles None/NaN.
    """
    # 1. Define Column Groups
    # numeric_cols = ["QUANTITE", "TONAGE", "RESTE T/P", "SURFACE"]
    # text_cols = ["NAVIRE", "B/L", "DESIGNATION", "CLIENT","DATE" ]
    # category_cols = ["TYPE", "SITUATION", "CLES"]

    cols_to_process = only_cols if only_cols is not None else datasource.columns
    for col in cols_to_process:
        if col not in datasource.columns:
            continue
        # --- Handle Text Columns (The "FLOAT" Error Fix) ---
        if col in text_cols:
            # Force to string first, then clean up the 'nan' strings
            datasource[col] = datasource[col].astype(str).replace(['nan', 'None', 'NaN', 'null'], '')
            datasource[col] = datasource[col].str.replace(r'\.0$', '', regex=True) # Removes .0 from IDs
            
        # --- Handle Numeric Columns ---
        elif col in numeric_cols:
            datasource[col] = pd.to_numeric(datasource[col], errors='coerce').fillna(0.0)
            if col == "QUANTITE":
                datasource[col] = datasource[col].astype(int)

        # --- Handle Date Columns ---
        elif col in date_cols:
            # Convert to datetime; errors='coerce' turns bad dates into NaT (Accepted by DateColumn)
            datasource[col] = pd.to_datetime(datasource[col], errors='coerce')

        # --- Handle Category Columns ---
        elif col in category_cols:
            default_val = "Divers" if col == "TYPE" else ("En attente" if col == "SITUATION" else "N/A")
            datasource[col] = datasource[col].astype(str).replace(['nan', 'None', 'NaN'], default_val)

    return datasource











def getDB():
    # 1. Check if the database file exists
    dir_name = os.path.dirname(DB_PATH)
    if not os.path.exists(dir_name):
        st.info(f"Database file creation at: {DB_PATH}")
        os.makedirs(dir_name)

    # 2. Create empty Excel file if it doesn't exist
    if not os.path.exists(DB_PATH):
        # Create a basic dataframe with columns
        df_new = pd.DataFrame(columns=COLUMNS)
        df_new.to_excel(DB_PATH, index=False)
        st.info(f"Created new database at: {DB_PATH}")

    # 2. Load the Master Data
    try:
        # We read from the single constant path now
        df = pd.read_excel(DB_PATH)
        return df
    except Exception as e:
        st.error(f"Error reading database: {e}")
        return null


def create_mapping_ui(uploaded_df, required_columns=COLUMNS):
    st.write("### Map Imported Columns to Database Columns")
    mapping = {}

    # Create a dropdown for every required column
    for req_col in required_columns:
        mapping[req_col] = st.selectbox(
            f"Select the source for: **{req_col}**",
            options=[None] + list(uploaded_df.columns),
            key=f"map_{req_col}"
        )
    return mapping


def align_data(uploaded_df, mapping):
    
    try:
        valid_mappings_count = sum(
            1 for value in mapping.values() if value is not None)

        if valid_mappings_count <= 2:
            return uploaded_df, False

        # Rename columns based on the mapping
        df_mapped = uploaded_df.rename(columns=mapping)

        final_cols = [value for key, value in mapping.items()
                      if value is not None]

        # Keep only the required columns
        df_aligned = df_mapped[final_cols]

        # Use KEYWORD_RULES from constants
        keyword_rules = KEYWORD_RULES

        # Ensure COL_DESIGNATION exists in the aligned DataFrame
        if COL_DESIGNATION in df_aligned.columns:

            def find_type(designation):
                """
                Return the CARGO TYPE by matching keywords against the designation.
                Rules are checked in order — first match wins (most specific first).
                
                KEYWORD_RULES format: [ ([keyword1, keyword2, ...], "TYPE"), ... ]
                """
                if not isinstance(designation, str):
                    return None

                designation_upper = designation.upper()

                # Iterate through rules in priority order
                for keywords, cargo_type in keyword_rules:
                    for keyword in keywords:
                        if keyword.upper() in designation_upper:
                            return cargo_type  # ← Returns the TYPE, not the keyword
                
                return None  # No match found → flag for manual review

            df_aligned[COL_TYPE] = df_aligned[COL_DESIGNATION].apply(find_type)

        else:
            # If COL_DESIGNATION is not present, set type column to None
            df_aligned[COL_TYPE] = df_aligned[COL_TYPE].fillna(value='None')

        return df_aligned, True

    except Exception as e:
        print(f"Error during alignment: {e}")
        return uploaded_df, False



@st.dialog("Map Your Columns", width="large")
def show_mapping_dialog(uploaded_df):
    st.write("Match your file columns to the database headings:")
    # st.info(list(uploaded_df.columns))
    st.session_state.trigger_mapping = False
    mapping = {}
    uploaded_cols = list(uploaded_df.columns)
    COLS_PER_ROW = 4

    for i in range(0, len(COLUMNS), COLS_PER_ROW):
        row_cols = st.columns(COLS_PER_ROW)
        batch = COLUMNS[i: i + COLS_PER_ROW]

        for j, req_col in enumerate(batch):
            with row_cols[j]:
                with st.container(border=True):
                    st.markdown(f"**{req_col}**")

                    # --- AUTO-MATCH LOGIC ---
                    # Find first uploaded col that contains the required name (e.g., 'date' in 'date_manifeste')
                    default_index = 0  # Default to None
                    for idx, col in enumerate(uploaded_cols):
                        if req_col.lower() in col.lower():
                            default_index = idx + 1 # +1 because [None] is at index 0
                            break
                    # ------------------------

                    selected_source_column = st.selectbox(
                        f"Source for {req_col}:",
                        options=[None] + uploaded_cols,
                        index=default_index,
                        key=f"map_{req_col}",
                        label_visibility="collapsed"
                    )
                    
                    if selected_source_column:
                        mapping[selected_source_column] = req_col
        
        st.session_state.final_mapping = mapping

    if st.button("Confirm and Import", type="primary", width='stretch'):
        # 1. Clear the trigger immediately so it doesn't re-open
        st.session_state.final_mapping = mapping
        # st.session_state.mapping_shown = True
        st.session_state.uploaded_file = None
        # Print to the terminal window

        # 2. Force a rerun to close the dialog and update the main app
        st.rerun()


def _compute_commodity_and_received_lines(raw_commodity: str, rec_str: str):
    """
    Given the raw commodity string from Excel and the received quantity string (rec_str),
    compute:
      - normalized commodity name
      - list of 'received' description lines
      - total_rec_str string to be displayed under 'Total Received'
    """
    # capital_comodity=raw_commodity.upper()
    commodity = raw_commodity
    received_lines = []
    total_rec_str = rec_str

    if matches_any_constant(raw_commodity, {"BAG", "BIG","BIG BAG","Calcined", "Anthracite", "Coal"}):
        commodity = "BIG BAGS"
        total_rec_str = f"{rec_str}  Big Bags"
        received_lines = [
            "BIG BAGS FOUND TORN ON BOARD",
            "BIG BAGS FOUND BROKEN ON BOARD",
            "EMPTY BAG ON BOARD",
        ]

    elif matches_any_constant(raw_commodity, {"PLYWOOD", "MDF", "CTP"}):
        commodity = raw_commodity.upper()
        received_lines = [
            f"Crates of {commodity} Found Dismembered on board",
            f"Crates of {commodity} wet on board (Packing and/or Contents)",
            f"Crates of {commodity} moldy on board (Packing and/or Contents)",
        ]
        total_rec_str = f"{rec_str}  Crates of {commodity}"

    elif matches_any_constant(raw_commodity, {"PIPE", "TUBE" }):
        commodity = "TUBES"
        received_lines = ["TUBES.", "TUBES Damaged on board"]
        total_rec_str = f"{rec_str}  {commodity}"

    elif matches_any_constant(raw_commodity, {"BEAMS"}):
        commodity = "Bundles of BEAMS"
        received_lines = ["Bundles of BEAMS.",
                          "Bundles of BEAMS Found Dismembered on board"]
        total_rec_str = f"{rec_str}  {commodity}"

    elif matches_any_constant(raw_commodity, {"formwork","steel moulds"}):
        commodity = "Bundles of formwork"
        received_lines = ["Bundles of formwork.",
                          "Bundles of formwork Found Dismembered on board"]
        total_rec_str = f"{rec_str}  {commodity}"
    
    elif matches_any_constant(raw_commodity, {"FILE MACHINE", "FIL","STEEL WIRE","WIRE","FIL M"}):
        commodity = "FIL MACHINE"
        received_lines = ["RLX FOUND DISMEMBERED ON BOARD"]
        total_rec_str = f"{rec_str}  {commodity}"

    elif matches_any_constant(raw_commodity, {"COIL", "BOB", "BOBINE"}):
        commodity = "COILS"
        received_lines = ["Coils Found Rusty on board",
                          "Coils Packaging damaged on board"]
        total_rec_str = f"{rec_str}  {commodity}"

    elif matches_any_constant(raw_commodity, {"WHITE WOOD", "BEECH WOOD", "RED WOOD"}):
        commodity = "BUNDLES"
        received_lines = [
            f"Bundles of {raw_commodity} Found Dismembered on board",
            f"Bundles of {raw_commodity} wet on board (Packing and/or Contents)",
            f"Bundles of {raw_commodity} moldy on board (Packing and/or Contents)",
        ]
        total_rec_str = f"{rec_str}  Bundles of {raw_commodity}"

    elif matches_any_constant(raw_commodity, {"COLI"}) and matches_any_constant(raw_commodity, {"PACKAGE"}):
        commodity = "Units + Package"
        received_lines = [commodity, f"{commodity} Damaged on board"]
        total_rec_str = f"{rec_str}  {commodity}"
    elif matches_any_constant(raw_commodity, {"TRACTEURS"
                "LOURD",
                "ENGINS",
                "UTILITAIRE"}):
        commodity = "Unit"
        received_lines = [commodity, f"{commodity} Damaged on board"]
        total_rec_str = f"{rec_str}  {commodity}"
    elif matches_any_constant(raw_commodity, {"PACKAGE"}):
        commodity = "package"
        received_lines = [commodity, f"{commodity} Damaged on board"]
        total_rec_str = f"{rec_str}  {commodity}"
    # elif not raw_commodity:
    #     commodity = "Units + Package"
    #     received_lines = ["Packaging damaged on board"]
    #     total_rec_str = f"{rec_str}  {commodity}"
    else:
        commodity = raw_commodity if raw_commodity else "General Cargo"
        received_lines = ["Packaging damaged on board"]
        total_rec_str = f"{rec_str}  {commodity}"
        
    return commodity, received_lines, total_rec_str


def _fill_entry_table(
    doc,
    table,
    client: str,
    commodity: str,
    manifest_qty_str: str,
    tonnage_str: str,
    received_lines,
    total_rec_str: str,
):
    """
    Fill the 5-row table in the DOCX document for a single cargo entry.
    Handles:
      - Receiver / Commodity
      - Manifested Quantity / Tonnage
      - Dynamic 'Received:' lines
      - Total Received line
      - Final note and separator
    """
    # Row 0: Receiver / Commodity
    row0 = table.rows[0].cells
    row0[0].width = Cm(9)
    p0 = row0[0].paragraphs[0]
    p0.add_run("Receiver : ").bold = True
    p0.add_run(client)
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT

    row0[1].width = Cm(9)
    p1 = row0[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_c = p1.add_run("Commodity : ")
    run_c1 = p1.add_run(commodity)
    run_c.bold = True
    run_c.font.name = "Agency FB"
    run_c1.font.name = "Agency FB"

    # Row 1: Manifested Quantity / Tonnage
    row1 = table.rows[1].cells
    row1[0].width = Cm(12)
    p2 = row1[0].paragraphs[0]
    p2.add_run("Manifested Quantity : ").bold = True
    p2.add_run(f"{manifest_qty_str} {commodity}")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    row1[1].width = Cm(5)
    p3 = row1[1].paragraphs[0]
    p3.add_run("Tonnage : ").bold = True
    p3.add_run(f"{tonnage_str} Mt")
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- DYNAMIC RECEIVED AREA (Row 2) ---
    row2 = table.rows[2].cells
    row2[0].width = Cm(30)

    row2_cell = table.rows[2].cells[0]
    # row2_cell.merge(table.rows[2].cells[1])  # if you want to merge later

    # Clear default paragraph and add the formatted lines
    row2_cell.paragraphs[0].clear()
    for i, line in enumerate(received_lines):
        if i == 0:
            p = row2_cell.paragraphs[0]
        else:
            p = row2_cell.add_paragraph()

        run_label = p.add_run("Received:    ")
        run_label.bold = True
        p.add_run(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 3: Total Received
    row3 = table.rows[3].cells
    row3[0].width = Cm(12)
    # row3[0].merge(row3[1])
    p4 = row3[0].paragraphs[0]
    p4.add_run("Total Received: ").bold = True
    p4.add_run(f" {total_rec_str}")
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 4: Final line
    row4 = table.rows[4].cells
    row4[0].width = Cm(25)
    # row4[0].merge(row4[1])
    p5 = row4[0].paragraphs[0]
    full = p5.add_run("The Quantity Will Be confirmed after delivery Cargo.")
    full.bold = True

    # Border Line
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = p_sep.add_run("=*"*29)
    run_sep.bold = True


def _shorten_bl_code(bl: str) -> str:
    """
    Shorten B/L codes like '25030TJD0701-17' or '25030TJD0101/01'
    Rule: Find 3 consecutive letters, then take those 3 letters and everything after them.
    Examples:
        '25030TJD0701-17' -> 'TJD0701-17'
        '25030TJD0101/01' -> 'TJD0101/01'
    """
    if bl is None:
        return ""
    s = str(bl).strip()

    # Find 3 consecutive letters
    m = re.search(r"([A-Za-z]{3})(.*)$", s)
    if m:
        # Return the 3 letters + everything after them
        return m.group(1) + m.group(2)

    return s

# Helper function to check if type matches any constant (partial/substring matching)


def matches_any_constant(type_str, constants_set):
    """
    Check if type_str contains any constant from constants_set (or vice versa).
    Handles partial matches like "ENGINS" matching "ENGIN" or "grue lourd" matching "GRUE".
    """
    type_str_upper = type_str.upper()
    # Check if any constant is a substring of the type, or type is a substring of constant
    for constant in constants_set:
        constant_upper = constant.upper()
        # Check if constant is contained in type (e.g., "ENGIN" in "ENGINS")
        if constant_upper in type_str_upper:
            return True
        # Check if type is contained in constant (e.g., "GRU" in "GRUE")
        if type_str_upper in constant_upper:
            return True
    return False




 # Fix: Use a helper function to avoid closure issues
def normalize_type(type_value):
    """Normalize TYPE value to standard categories"""
    if pd.isna(type_value):
        return None
    
    type_str = str(type_value).strip().upper()
    
    if matches_any_constant(type_str, UNITS_TYPES):
        return "UNITS"
    elif matches_any_constant(type_str, PACKAGES_TYPES):
        return "PACKAGES"
    else:
        return type_str  # Keep original if not matching




def first_non_null(series):
    return next((x for x in series if pd.notna(x)), None)


# Helper function for B/L aggregation
def aggregate_bl(series):
    bl_values = [str(x).strip()
                 for x in series if pd.notna(x) and str(x).strip()]
    if not bl_values:
        return ""
    unique_bls = list(pd.Series(bl_values).unique())
    return ",".join(unique_bls)


def group_sourcefile_by_client(
    input_excel: str,
    sheet_name: int | str = 0,
    skip_units_packages: bool = False,
    bl_aggregated: bool = False,

) -> pd.DataFrame:
    df = pd.read_excel(input_excel, sheet_name=sheet_name, engine="openpyxl")

      # Skip rows whose commodity type is in UNITS_TYPES or PACKAGES_TYPES
    if skip_units_packages and COL_TYPE in df.columns:
        skip_types = UNITS_TYPES | PACKAGES_TYPES
        df = df[
            ~df[COL_TYPE]
            .astype(str)
            .str.strip()
            .str.upper()
            .isin(skip_types)
        ].reset_index(drop=True)

    # Ensure numeric
    for col in [COL_QUANTITE, COL_TONAGE]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)

    # Short B/L into the same column
    if COL_BL in df.columns:
        df[COL_BL] = df[COL_BL].apply(_shorten_bl_code)
    

        # ============================================================
    # NEW: Normalize TYPE column before grouping
    # ============================================================
    if COL_TYPE in df.columns:
        df[COL_TYPE] = df[COL_TYPE].apply(normalize_type)


    # Base aggregation
    agg_dict = {
        COL_QUANTITE: "sum",
        COL_TONAGE: "sum",
        COL_BL:aggregate_bl,
    }

    skip_cols=[COL_CLIENT, COL_QUANTITE, COL_TONAGE,COL_BL,COL_TYPE]
    
    if not bl_aggregated  :
        skip_cols.remove(COL_BL)
        agg_dict.pop(COL_BL, None)

    # For all other columns, keep first non-null value
    for col in COLUMNS:
        if col in skip_cols:
            continue
        
        if col in df.columns:
            agg_dict[col] = first_non_null
           
    grouped = df.groupby([COL_CLIENT, COL_TYPE], as_index=False).agg(agg_dict)

    sorted_grouped = grouped.sort_values(
        by=COL_TYPE, ascending=True, na_position='last').reset_index(drop=True)

    return sorted_grouped
