import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from assets.constants.constants import PATH_PVS


def generate_daily_pv(file_path):
    # --- 1. SETUP PATHS & FILENAME ---
    base_name = os.path.basename(file_path)
    file_name_only = os.path.splitext(base_name)[0]
    

    # Create directory named after the source file
    output_dir =f"{PATH_PVS}/{file_name_only}"
    # os.path.join(os.getcwd(), file_name_only)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    current_date = datetime.now().strftime("%d-%m-%Y")
    output_filename = f"PV DEB GROUPE B {current_date}.docx"
    save_path = os.path.join(output_dir, output_filename)

    # --- 2. LOAD DATA ---
    # Assuming Table 1 is in Sheet 0 and Table 2 is in Sheet 1 
    # Or adjust if all data is in one sheet
    df_cargo = pd.read_excel(file_path, sheet_name=0) 
    df_pointers = pd.read_excel(file_path, sheet_name=0) # Change if needed
    
    # Mock variables (Replace with actual logic to extract from Excel if needed)
    ship_name = file_name_only 
    report_date = current_date

    # --- 3. CREATE DOCUMENT ---
    doc = Document()

    # 1. Date
    date_paragraph = doc.add_paragraph(report_date)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 2. Title
    title = doc.add_heading('PV de debarquement', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Intro
    doc.add_paragraph(f"Navire {ship_name} du {report_date}, on a conclus ce qui suit:")

    # 4. Cargo Table Text
    doc.add_paragraph("voici l'état de débarquement de la marchandise:", style='List Bullet')

    # 5. Table 1 (Cargo)
    table1 = doc.add_table(rows=len(df_cargo) + 1, cols=5)
    table1.style = 'Table Grid'
    
    headers = ['ID', 'Description', 'Poids', 'Zone', 'Etat']
    for i, h in enumerate(headers):
        table1.rows[0].cells[i].text = h

    for i in range(len(df_cargo)):
        row_cells = table1.rows[i+1].cells
        for j in range(5):
            # Safe check to avoid index errors if Excel has fewer columns
            val = df_cargo.iloc[i, j] if j < len(df_cargo.columns) else ""
            row_cells[j].text = str(val)

    doc.add_paragraph() 

    # 6. Conclusion & Pointers
    p_concl = doc.add_paragraph()
    p_concl.add_run("Conclusion: ").bold = True
    doc.add_paragraph("Liste des pointeurs du groupe « B »:", style='List Bullet')

    # 7. Table 2 (Group B)
    table2 = doc.add_table(rows=len(df_pointers) + 2, cols=2)
    table2.style = 'Table Grid'

    # Header Merge
    row1_cells = table2.rows[0].cells
    merged = row1_cells[0].merge(row1_cells[1])
    merged.text = "GROUPE B"
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sub-Headers
    table2.rows[1].cells[0].text = "MV"
    table2.rows[1].cells[1].text = "Reçu TP"

    # Fill Data
    for i in range(len(df_pointers)):
        row_cells = table2.rows[i+2].cells
        row_cells[0].text = str(df_pointers.iloc[i, 0])
        row_cells[1].text = str(df_pointers.iloc[i, 1])

    # --- 4. SAVE ---
    doc.save(save_path)
    return save_path