import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# STEP 1: LOAD DATA (Simulating your Excel)
# ==========================================
# In a real scenario, you would use: df = pd.read_excel('your_file.xlsx')

# ==========================================
# STEP 2: CREATE THE DOCUMENT
# ==========================================

doc = Document()

# --- 1. Date (Aligned Right) ---
date_paragraph = doc.add_paragraph(report_date)
date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# --- 2. Static Title ---
title = doc.add_heading('PV de debarquement', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 3. Intro Sentence ---
intro_text = f"Navire {ship_name} du {report_date} ,on a conclus ce qui suit:"
doc.add_paragraph(intro_text)

# --- 4. Bullet Point ---
# 'List Bullet' is a standard Word style
doc.add_paragraph(
    "voici l'état de débarquement de la marchandise dans le tableaux suivant :", 
    style='List Bullet'
)

# --- 5. Table 1: Cargo (5 Columns) ---
# Add table with 1 header row + rows from dataframe
row_count = len(df_cargo) + 1 
table1 = doc.add_table(rows=row_count, cols=5)
table1.style = 'Table Grid' # Adds borders

# Add Headers manually (or from dataframe columns)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Poids'
hdr_cells[3].text = 'Zone'
hdr_cells[4].text = 'Etat'

# Fill the table with data
for i in range(len(df_cargo)):
    row_cells = table1.rows[i+1].cells
    # Loop through columns (0 to 4)
    for j in range(5):
        row_cells[j].text = str(df_cargo.iloc[i, j])

doc.add_paragraph() # Add a little space

# --- 6. Conclusion Section ---
p_concl = doc.add_paragraph()
runner = p_concl.add_run("Conclusion: ")
runner.bold = True

doc.add_paragraph(
    "Voici la liste des pointeurs du groupe « B » affecter au navire:",
    style='List Bullet'
)

# --- 7. Table 2: Group B (Merged Header) ---
# We need 2 base columns. 
# Rows needed: 1 (Merged Title) + 1 (Headers MV/TP) + Data Rows
total_rows_t2 = len(df_pointers) + 2
table2 = doc.add_table(rows=total_rows_t2, cols=2)
table2.style = 'Table Grid'

# Row 1: Merge Cells and add "GROUPE B"
row1_cells = table2.rows[0].cells
merged_cell = row1_cells[0].merge(row1_cells[1])
merged_cell.text = "GROUPE B"
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Row 2: Headers "MV" and "Reçu TP"
row2_cells = table2.rows[1].cells
row2_cells[0].text = "MV"
row2_cells[1].text = "Reçu TP"

# Subsequent Rows: Data from Excel
for i in range(len(df_pointers)):
    row_cells = table2.rows[i+2].cells
    row_cells[0].text = str(df_pointers.iloc[i, 0]) # MV Column
    row_cells[1].text = str(df_pointers.iloc[i, 1]) # Reçu TP Column

# ==========================================
# STEP 3: SAVE
# ==========================================
doc.save('PV_Debarquement.docx')
print("Document created successfully.")

















import os
from datetime import datetime

def generate_daily_pv(file_path):
    # 1. Get file name and create directory
    base_name = os.path.basename(file_path)
    file_name_only = os.path.splitext(base_name)[0]
    
    # Create folder with the same name as the Excel file
    output_dir = os.path.join(os.getcwd(), file_name_only)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 2. Prepare Date and Filename
    current_date = datetime.now().strftime("%d-%m-%Y")
    output_filename = f"PV DEB GROUPE B {current_date}.docx"
    save_path = os.path.join(output_dir, output_filename)
    
    # 3. Load Data
    df = pd.read_excel(file_path)
    
    # 4. Your Word Generation Logic (Simplified for brevity)
    doc = Document()
    # ... (Add your title, tables, and logic from previous code here) ...
    
    # 5. Save
    doc.save(save_path)
    return save_path