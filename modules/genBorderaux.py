import os
import math
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from assets.constants.constants import (
    PATH_BRDX,
    PATH_TEMPLATES,
    COL_CLIENT,
    COL_TYPE,
    COL_QUANTITE,
    COL_TONAGE,
    COL_RESTE_TP,
    # add more if you need them
)

import streamlit as st

from tools.tools import _compute_commodity_and_received_lines, _fill_entry_table, group_sourcefile_by_client

if not os.path.exists(PATH_BRDX):
    os.makedirs(PATH_BRDX)


def clean_excel_val(val):
    # Check for None or empty string
    if val is None or val == "":
        return 0
    # Check for NaN (only if val is a number)
    if isinstance(val, (int, float)) and math.isnan(val):
        return 0
    return val


def format_entry_docx(doc, row):
    client = str(row.get(COL_CLIENT, "")).strip()
    # Initial commodity from excel
    raw_commodity = str(row.get(COL_TYPE, "")).strip().upper()
    
    # nb_colis =   0     if pd.notna(row.get("nb_colis")) else  row.get("nb_colis")
    # tonnage  =   0.0   if pd.notna(row.get("tonnage"))  else  row.get("tonnage")
    # rec_qty  =0 if pd.notna(row.get("rec_qty"))  else  row.get("rec_qty")

    # print(f"client------{client}")
    nb_colis = row.get(COL_QUANTITE)
    tonnage = row.get(COL_TONAGE)
    rec_qty = row.get(COL_RESTE_TP)or 0
    
    

    nb_colis = clean_excel_val(nb_colis)
    tonnage = clean_excel_val(tonnage)
    rec_qty = clean_excel_val(rec_qty)

    # Create table
    table = doc.add_table(rows=5, cols=2)
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tonnage_str = f"{tonnage:.2f}".lstrip(
        "0") if tonnage < 1 else f"{tonnage:.2f}"
    manifest_qty_str = f"{int(nb_colis):02d}"
    rec_str = f"{int(rec_qty):02d}"

    # Define lines based on type
    # Logic to adjust Commodity name based on type
    # Initialize defaults
    commodity, received_lines, total_rec_str = _compute_commodity_and_received_lines(
        raw_commodity,
        rec_str,
    )

    # --- use helper to fill the table and separator in the document ---
    _fill_entry_table(
        doc=doc,
        table=table,
        client=client,
        commodity=commodity,
        manifest_qty_str=manifest_qty_str,
        tonnage_str=tonnage_str,
        received_lines=received_lines,
        total_rec_str=total_rec_str,
    )


def excel_to_docx_custom(input_excel, sheet_name=0, template_path=None, output_docx=None):
    if not output_docx:
        return
    
    # Accept either a path or a DataFrame
    if isinstance(input_excel, str):
        df = pd.read_excel(input_excel, sheet_name=sheet_name,
                           engine="openpyxl", header=0)
    else:
        df = input_excel  # already a DataFrame

    # df = pd.read_excel(input_excel, sheet_name=sheet_name,
    #                    engine="openpyxl", header=0)

    doc = Document(template_path) if template_path else Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri (Corps)"
    font.size = Pt(12)

    for idx, row in df.iterrows():
        format_entry_docx(doc, row)

    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    if os.path.exists(output_docx):
        os.remove(output_docx)
        # print(f"{output_docx} has been deleted.")

    doc.save(output_docx)
    # print(f"New File {output_docx} Saved")


def generate_brd(sourcefile, sheet_name=0, template_name="template.docx"):
    base_name = os.path.basename(sourcefile)
    file_name_only = os.path.splitext(base_name)[0]
    output_docx = f"{PATH_BRDX}/{file_name_only}.docx"
    template_path = f"{PATH_TEMPLATES}/{template_name}"

    grouped_df = group_sourcefile_by_client(sourcefile)
    
    st.dataframe(grouped_df)      # nicer interactive table
    excel_to_docx_custom(grouped_df, sheet_name, template_path, output_docx)
    return output_docx

# if __name__ == "__main__":
    # Ensure book1.xlsx exists in your directory
    # excel_to_docx_custom("book1.xlsx", output_docx="entries.docx")
    # column_names = ["type", "client", "qte", "poids", "rec_qty"]
    # names=column_names
#    generate_brd("source.xlsx", sheet_name=0, template_path="template.docx", output_docx="entries.docx")

